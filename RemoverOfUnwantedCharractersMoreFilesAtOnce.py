import docx
import tkinter as tk
from tkinter import filedialog
import os
import re

def remove_unwanted_numbers(file_paths):
    # Vytvorenie nového dokumentu pre všetky spracované súbory
    combined_doc = docx.Document()
    
    # Prechádzanie všetkých vybraných súborov v opačnom poradí (od najstaršieho po najnovší)
    for file_path in reversed(file_paths):
        # Načítanie dokumentu Word
        doc = docx.Document(file_path)
        
        # Prechádzanie všetkých odsekov v dokumente
        for paragraph in doc.paragraphs:
            # Získanie textu odseku
            text = paragraph.text
            
            # Pravidlo na odstránenie čísel, ktoré sú obklopené medzerami, ale zachovanie čísel kapitol a strán
            new_text = re.sub(r'\b\d{4,}\b', '', text).strip()  # Odstráni čísla s 4 a viac číslicami
            
            # Pridanie upraveného textu do kombinovaného dokumentu
            combined_doc.add_paragraph(new_text)
    
    # Získanie názvu prvého vstupného súboru bez prípony
    base_name = os.path.splitext(os.path.basename(file_paths[0]))[0]
    
    # Vytvorenie názvu výstupného súboru
    output_path = f"{base_name}AllBookCleaned.docx"
    
    # Uloženie nového dokumentu
    combined_doc.save(output_path)
    
    # Vypísať absolútnu cestu k výstupnému súboru
    print(f"Číselné znaky boli úspešne odstránené. Výstupný súbor: {os.path.abspath(output_path)}")

def open_file_dialog():
    # Otvorenie dialógového okna na výber súborov
    root = tk.Tk()
    root.withdraw()  # Skrytie hlavného okna
    file_paths = filedialog.askopenfilenames(title="Vyberte Word súbor/y", filetypes=[("Word súbory", "*.docx")])
    
    if file_paths:  # Ak boli vybrané súbory
        remove_unwanted_numbers(file_paths)
    else:
        print("Nebol vybraný žiaden súbor.")

# Spustenie dialógového okna
open_file_dialog()