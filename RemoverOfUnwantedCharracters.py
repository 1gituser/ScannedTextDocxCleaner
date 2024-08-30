import docx
import tkinter as tk
from tkinter import filedialog
import os
import re

def remove_unwanted_numbers(file_path):
    # Načítanie dokumentu Word
    doc = docx.Document(file_path)
    
    # Vytvorenie nového dokumentu
    new_doc = docx.Document()
    
    # Prechádzanie všetkých odsekov v dokumente
    for paragraph in doc.paragraphs:
        # Získanie textu odseku
        text = paragraph.text
        
        # Pravidlo na odstránenie čísel, ktoré sú obklopené medzerami, ale zachovanie čísel kapitol a strán
        # Zachovanie čísel s maximálne 3 číslicami
        new_text = re.sub(r'\b\d{4,}\b', '', text).strip()  # Odstráni čísla s 4 a viac číslicami
        
        # Pridanie upraveného textu do nového dokumentu
        new_doc.add_paragraph(new_text)
    
    # Získanie názvu vstupného súboru bez prípony
    base_name = os.path.splitext(os.path.basename(file_path))[0]
    
    # Vytvorenie názvu výstupného súboru
    output_path = f"{base_name}Cleaned.docx"
    
    # Uloženie nového dokumentu
    new_doc.save(output_path)
    
    # Vypísať absolútnu cestu k výstupnému súboru
    print(f"Číselné znaky boli úspešne odstránené. Výstupný súbor: {os.path.abspath(output_path)}")

def open_file_dialog():
    # Otvorenie dialógového okna na výber súboru
    root = tk.Tk()
    root.withdraw()  # Skrytie hlavného okna
    file_path = filedialog.askopenfilename(title="Vyberte Word súbor", filetypes=[("Word súbory", "*.docx")])
    
    if file_path:  # Ak bol vybraný súbor
        remove_unwanted_numbers(file_path)
    else:
        print("Nebolo vybrané žiadne súbor.")

# Spustenie dialógového okna
open_file_dialog()